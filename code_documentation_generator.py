import os
import time
import re
import json
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from dotenv import load_dotenv
from azure.ai.projects import AIProjectClient
from azure.identity import DefaultAzureCredential
from json_repair import repair_json

class CodeDocumentationGenerator:
    """
    A class that generates technical and business documentation for a codebase
    using Azure AI services.
    """
    
    def __init__(self, conn_str=None, agent_name=None, model_id=None, code_directory=None, env_file=None):
        """
        Initialize the CodeDocumentationGenerator with required parameters.
        
        Args:
            conn_str (str, optional): Azure connection string
            agent_name (str, optional): Name of the AI agent
            model_id (str, optional): ID of the model to use
            code_directory (str, optional): Directory containing the code to document
            env_file (str, optional): Path to .env file to load environment variables from
        """
        # Load environment variables if env_file is provided
        if env_file and os.path.exists(env_file):
            load_dotenv(env_file)
            
        # Use provided parameters or environment variables
        self.conn_str = conn_str or os.getenv("CONN_STR")
        self.agent_name = agent_name or os.getenv("DOC_AGENT_NAME")
        self.model_id = model_id or os.getenv("MODEL_ID")
        self.code_directory = code_directory or os.getenv("CODE_DIRECTORY")
        self.output_dir = os.getenv("OUTPUT_DIR")
        # Initialize client and other properties
        self.project_client = None
        self.agent = None
        self.thread = None
        
        # Validate required parameters
        if not all([self.conn_str, self.agent_name, self.model_id, self.code_directory]):
            missing = []
            if not self.conn_str: missing.append("conn_str")
            if not self.agent_name: missing.append("agent_name")
            if not self.model_id: missing.append("model_id")
            if not self.code_directory: missing.append("code_directory")
            raise ValueError(f"Missing required parameters: {', '.join(missing)}")
    
    def initialize_client(self):
        """Initialize the Azure AI Project client."""
        try:
            self.project_client = AIProjectClient.from_connection_string(
                credential=DefaultAzureCredential(),
                conn_str=self.conn_str,
            )
            print("✅ Client initialized successfully.")
            return True
        except Exception as e:
            print(f"❌ Error initializing client: {e}")
            return False
    
    def setup_agent(self):
        """Check if the agent exists, create it if not."""
        try:
            # Try to find the agent by name
            agents = self.project_client.agents.list_agents()
            self.agent = None
            
            for existing_agent in agents.data:
                if existing_agent.name == self.agent_name:
                    self.agent = existing_agent
                    print(f"✅ Agent found: {self.agent.name}")
                    break
    
            if self.agent is None:
                # If the agent is not found, create it
                self.agent = self.project_client.agents.create_agent(
                    model=self.model_id,
                    name=self.agent_name,
                    description="An agent that analyzes code and generates documentation",
                    instructions="You are an AI agent that analyzes code and generates both technical and business documentation.",
                    temperature=0.01,
                )
                print(f"✅ Agent created: {self.agent.name}")
            
            return True
        except Exception as e:
            print(f"❌ Error checking/creating agent: {e}")
            return False
    
    def create_thread(self):
        """Create a communication thread."""
        try:
            self.thread = self.project_client.agents.create_thread()
            print("✅ Thread created.")
            return True
        except Exception as e:
            print(f"❌ Error creating thread: {e}")
            return False
    
    def collect_code_files(self):
        """
        Read code from the specified directory and prepare a summary.
        
        Returns:
            str: A summary of all code files
        """
        code_dir = self.code_directory
        code_summary = ""
        
        for root, dirs, files in os.walk(code_dir):
            for file in files:
                file_path = os.path.join(root, file)
                if file.endswith((".py", ".md", ".txt", ".yaml")) or file in ("Dockerfile", ".env.example"):
                    try:
                        with open(file_path, "r", encoding="utf-8") as f:
                            content = f.read()
                            relative_path = os.path.relpath(file_path, code_dir)
                            file_ext = file.split('.')[-1] if '.' in file else ''
                            code_summary += f"\n\n### File: {relative_path}\n```{file_ext}\n{content}\n```\n"
                    except Exception as e:
                        print(f"⚠️ Failed to read {file_path}: {e}")
        
        return code_summary
    
    def send_message(self, code_summary):
        """
        Compose and send a message to the agent.
        
        Args:
            code_summary (str): Summary of all code files
            
        Returns:
            bool: True if successful, False otherwise
        """
        user_message = (
                    "You are an AI agent that generates structured documentation for a codebase.\n"
                    "Please analyze the provided code and return documentation in the following JSON format:\n"
                    "{\n"
                    '  "technical": "<technical documentation markdown>",\n'
                    '  "business": "<business documentation markdown>"\n'
                    "}\n"
                    f"{code_summary}"
                )
        try:
            self.project_client.agents.create_message(
                thread_id=self.thread.id,
                role="user",
                content=user_message
            )
            print("✅ Message sent.")
            return True
        except Exception as e:
            print(f"❌ Error sending message: {e}")
            return False
    
    def run_agent(self):
        """Run the agent on the thread."""
        try:
            run = self.project_client.agents.create_and_process_run(
                thread_id=self.thread.id,
                agent_id=self.agent.id
            )
            print("✅ Agent run started.")
            
            # Wait a moment for processing to begin
            time.sleep(0.5)
            return True
        except Exception as e:
            print(f"❌ Error processing run: {e}")
            return False
    
    def retrieve_responses(self):
        """
        Retrieve and process responses from the agent.
        
        Returns:
            tuple: (technical_doc, business_doc) or (None, None) if failed
        """
        try:
            messages = self.project_client.agents.list_messages(thread_id=self.thread.id)
            print("✅ Messages retrieved.")
            
            if not hasattr(messages, "data") or not messages.data:
                print("⚠️ No messages found.")
                return None, None
                
            sorted_messages = sorted(messages.data, key=lambda x: x.created_at)
            full_response = ""
            
            for msg in sorted_messages:
                if msg.content and isinstance(msg.content, list):
                    for content_item in msg.content:
                        if content_item["type"] == "text":
                            full_response += content_item["text"]["value"] + "\n"
            
            print("✅ Full response constructed.")
            print(full_response)
            ## Extract the JSON block using regex to ensure we catch a proper JSON object
            try:
                json_block_match = re.search(r"```json\s*(\{.*?\})\s*```", full_response, re.DOTALL)
                if json_block_match:
                    # Extract the JSON block from the match
                    json_response = json.loads(repair_json(json_block_match.group(1)))  # Use the captured group, not the full match
                    print(json.dumps(json_response))
                else:
                    json_response = json.loads(repair_json(full_response[full_response.lower().find("{"):].strip()))
                    print(json.dumps(json_response))
            except json.JSONDecodeError as e:
                print(f"⚠️ JSON parsing error: {e}")
                return None, None
            if isinstance(json_response, list):
                json_response = json_response[0]
            tech_doc = json_response.get("technical", "").strip()
            biz_doc = json_response.get("business", "").strip()
            
            if tech_doc and biz_doc:
                return tech_doc, biz_doc
            else:
                print("⚠️ JSON keys found but content is empty or missing.")
                return None, None
            
        except Exception as e:
            print(f"❌ Error retrieving messages: {e}")
            return None, None
    
    def save_documentation(self, tech_doc, biz_doc):
        """
        Save technical and business documentation to files.
        
        Args:
            tech_doc (str): Technical documentation content
            biz_doc (str): Business documentation content
            output_dir (str, optional): Directory to save documentation files
                                        (defaults to code_directory)
        
        Returns:
            bool: True if successful, False otherwise
        """
        if not tech_doc or not biz_doc:
            print("⚠️ Cannot save empty documentation.")
            return False
            
        output_directory = self.output_dir
        
        try:
            # Save markdown files
            with open(os.path.join(output_directory, "README_TECHNICAL.md"), "w", encoding="utf-8") as f:
                f.write("# Technical Documentation\n\n" + tech_doc)
            with open(os.path.join(output_directory, "README_BUSINESS.md"), "w", encoding="utf-8") as f:
                f.write("# Business Documentation\n\n" + biz_doc)
            print("✅ Documentation files created successfully.")
            
            # Convert to DOCX with proper formatting
            self.convert_markdown_to_docx(
                tech_doc,
                os.path.join(output_directory, "README_TECHNICAL.docx")
            )
            self.convert_markdown_to_docx(
                biz_doc,
                os.path.join(output_directory, "README_BUSINESS.docx")
            )
            print("✅ DOCX documentation files created successfully.")
            
            return True
        except Exception as e:
            print(f"❌ Error saving documentation: {e}")
            return False
    
    @staticmethod
    def convert_markdown_to_docx(markdown_text, output_file):
        """
        Convert markdown text to a Word document with proper formatting.
        
        Args:
            markdown_text (str): The markdown text to convert
            output_file (str): Path to save the output docx file
        """
        # Create a new Document
        doc = Document()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(
            markdown_text,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )
        
        # Parse the HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Process each element
        for element in soup.find_all():
            # Process headings
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Get heading level
                level = int(element.name[1])
                heading = doc.add_heading('', level=level)
                heading.add_run(element.get_text()).bold = True
                
            # Process paragraphs
            elif element.name == 'p':
                if element.find('code'):
                    # Inline code in paragraph
                    p = doc.add_paragraph()
                    for content in element.contents:
                        if hasattr(content, 'name') and content.name == 'code':
                            run = p.add_run(content.get_text())
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                        else:
                            p.add_run(str(content))
                else:
                    # Regular paragraph
                    doc.add_paragraph(element.get_text())
                    
            # Process code blocks
            elif element.name == 'pre':
                # Code block
                code_text = element.get_text()
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                code_run = p.add_run(code_text)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.5)
                
            # Process lists
            elif element.name == 'li':
                # Check if parent is ul or ol
                if element.parent.name == 'ul':
                    # Unordered list item
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(element.get_text())
                elif element.parent.name == 'ol':
                    # Ordered list item
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(element.get_text())
                    
            # Process horizontal rules
            elif element.name == 'hr':
                # Horizontal rule
                doc.add_paragraph('_' * 50)
        
        # Save the document
        doc.save(output_file)
    
    def generate_documentation(self, output_dir=None):
        """
        Generate documentation for the codebase in one method call.
        
        Args:
            output_dir (str, optional): Directory to save documentation files
                                         (defaults to code_directory)
        
        Returns:
            bool: True if documentation was generated successfully, False otherwise
        """
        if not self.initialize_client():
            return False
            
        if not self.setup_agent():
            return False
            
        if not self.create_thread():
            return False
            
        code_summary = self.collect_code_files()
        if not code_summary:
            print("⚠️ No code files found to document.")
            return False
            
        if not self.send_message(code_summary):
            return False
            
        if not self.run_agent():
            return False
            
        tech_doc, biz_doc = self.retrieve_responses()
        if not tech_doc or not biz_doc:
            return False
            
        return self.save_documentation(tech_doc, biz_doc)
